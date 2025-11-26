"""
Export Service for MultinotesAI.

This module provides export functionality for:
- PDF export with formatting
- DOCX export with styles
- Markdown export
- Plain text export
- JSON export

WBS Items:
- 4.4.9: Build export functionality (PDF, DOCX, MD)
- 6.2.1-6.2.3: Export formats
"""

import io
import json
import logging
import re
from datetime import datetime
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass
from enum import Enum

from django.conf import settings
from django.template.loader import render_to_string
from django.utils import timezone

logger = logging.getLogger(__name__)


# =============================================================================
# Export Types
# =============================================================================

class ExportFormat(Enum):
    """Supported export formats."""
    PDF = 'pdf'
    DOCX = 'docx'
    MARKDOWN = 'md'
    TEXT = 'txt'
    JSON = 'json'
    HTML = 'html'


@dataclass
class ExportContent:
    """Content to be exported."""
    title: str
    content: str
    created_at: Optional[datetime] = None
    updated_at: Optional[datetime] = None
    author: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None


@dataclass
class ConversationExport:
    """Conversation data for export."""
    title: str
    messages: List[Dict[str, str]]  # [{'role': 'user/assistant', 'content': '...'}]
    model_name: Optional[str] = None
    created_at: Optional[datetime] = None
    total_tokens: Optional[int] = None


# =============================================================================
# Markdown Exporter
# =============================================================================

class MarkdownExporter:
    """Export content to Markdown format."""

    def export_content(self, content: ExportContent) -> str:
        """
        Export single content to Markdown.

        Args:
            content: Content to export

        Returns:
            Markdown formatted string
        """
        lines = []

        # Title
        lines.append(f"# {content.title}\n")

        # Metadata
        if content.author:
            lines.append(f"**Author:** {content.author}\n")
        if content.created_at:
            lines.append(f"**Created:** {content.created_at.strftime('%Y-%m-%d %H:%M')}\n")
        if content.updated_at:
            lines.append(f"**Updated:** {content.updated_at.strftime('%Y-%m-%d %H:%M')}\n")

        lines.append("\n---\n")

        # Content
        lines.append(content.content)

        return "\n".join(lines)

    def export_conversation(self, conversation: ConversationExport) -> str:
        """
        Export conversation to Markdown.

        Args:
            conversation: Conversation to export

        Returns:
            Markdown formatted string
        """
        lines = []

        # Header
        lines.append(f"# {conversation.title}\n")

        if conversation.model_name:
            lines.append(f"**Model:** {conversation.model_name}\n")
        if conversation.created_at:
            lines.append(f"**Date:** {conversation.created_at.strftime('%Y-%m-%d %H:%M')}\n")
        if conversation.total_tokens:
            lines.append(f"**Tokens Used:** {conversation.total_tokens:,}\n")

        lines.append("\n---\n")

        # Messages
        for msg in conversation.messages:
            role = msg.get('role', 'user')
            content = msg.get('content', '')

            if role == 'user':
                lines.append(f"## 👤 User\n")
            elif role == 'assistant':
                lines.append(f"## 🤖 Assistant\n")
            else:
                lines.append(f"## {role.title()}\n")

            lines.append(f"{content}\n")
            lines.append("\n---\n")

        return "\n".join(lines)

    def export_multiple(self, contents: List[ExportContent]) -> str:
        """
        Export multiple contents to a single Markdown document.

        Args:
            contents: List of contents to export

        Returns:
            Markdown formatted string
        """
        sections = []
        for i, content in enumerate(contents, 1):
            sections.append(f"# {i}. {content.title}\n")
            sections.append(content.content)
            sections.append("\n---\n")

        return "\n".join(sections)


# =============================================================================
# Plain Text Exporter
# =============================================================================

class TextExporter:
    """Export content to plain text format."""

    def export_content(self, content: ExportContent) -> str:
        """Export content to plain text."""
        lines = []

        lines.append(content.title.upper())
        lines.append("=" * len(content.title))
        lines.append("")

        if content.author:
            lines.append(f"Author: {content.author}")
        if content.created_at:
            lines.append(f"Created: {content.created_at.strftime('%Y-%m-%d %H:%M')}")

        lines.append("")
        lines.append("-" * 40)
        lines.append("")

        # Strip HTML/Markdown formatting
        clean_content = self._strip_formatting(content.content)
        lines.append(clean_content)

        return "\n".join(lines)

    def export_conversation(self, conversation: ConversationExport) -> str:
        """Export conversation to plain text."""
        lines = []

        lines.append(conversation.title.upper())
        lines.append("=" * len(conversation.title))
        lines.append("")

        if conversation.model_name:
            lines.append(f"Model: {conversation.model_name}")
        if conversation.created_at:
            lines.append(f"Date: {conversation.created_at.strftime('%Y-%m-%d %H:%M')}")

        lines.append("")
        lines.append("-" * 40)

        for msg in conversation.messages:
            role = msg.get('role', 'user').upper()
            content = self._strip_formatting(msg.get('content', ''))

            lines.append("")
            lines.append(f"[{role}]")
            lines.append(content)
            lines.append("-" * 40)

        return "\n".join(lines)

    def _strip_formatting(self, text: str) -> str:
        """Remove HTML/Markdown formatting from text."""
        # Remove HTML tags
        text = re.sub(r'<[^>]+>', '', text)

        # Remove Markdown formatting
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Bold
        text = re.sub(r'\*(.+?)\*', r'\1', text)  # Italic
        text = re.sub(r'__(.+?)__', r'\1', text)  # Bold alt
        text = re.sub(r'_(.+?)_', r'\1', text)  # Italic alt
        text = re.sub(r'`(.+?)`', r'\1', text)  # Code
        text = re.sub(r'#{1,6}\s*', '', text)  # Headers
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)  # Links

        return text


# =============================================================================
# JSON Exporter
# =============================================================================

class JSONExporter:
    """Export content to JSON format."""

    def export_content(self, content: ExportContent) -> str:
        """Export content to JSON."""
        data = {
            'title': content.title,
            'content': content.content,
            'created_at': content.created_at.isoformat() if content.created_at else None,
            'updated_at': content.updated_at.isoformat() if content.updated_at else None,
            'author': content.author,
            'metadata': content.metadata or {},
            'exported_at': timezone.now().isoformat(),
        }
        return json.dumps(data, indent=2, ensure_ascii=False)

    def export_conversation(self, conversation: ConversationExport) -> str:
        """Export conversation to JSON."""
        data = {
            'title': conversation.title,
            'model': conversation.model_name,
            'created_at': conversation.created_at.isoformat() if conversation.created_at else None,
            'total_tokens': conversation.total_tokens,
            'messages': conversation.messages,
            'exported_at': timezone.now().isoformat(),
        }
        return json.dumps(data, indent=2, ensure_ascii=False)

    def export_multiple(self, contents: List[ExportContent]) -> str:
        """Export multiple contents to JSON array."""
        data = []
        for content in contents:
            data.append({
                'title': content.title,
                'content': content.content,
                'created_at': content.created_at.isoformat() if content.created_at else None,
                'author': content.author,
            })
        return json.dumps(data, indent=2, ensure_ascii=False)


# =============================================================================
# HTML Exporter
# =============================================================================

class HTMLExporter:
    """Export content to HTML format."""

    def __init__(self):
        self.template_name = 'exports/content.html'
        self.conversation_template = 'exports/conversation.html'

    def export_content(self, content: ExportContent) -> str:
        """Export content to HTML."""
        # Convert Markdown to HTML if needed
        html_content = self._markdown_to_html(content.content)

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{content.title}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 2rem;
            line-height: 1.6;
            color: #333;
        }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 0.5rem; }}
        .meta {{ color: #7f8c8d; font-size: 0.9rem; margin-bottom: 1rem; }}
        .content {{ margin-top: 2rem; }}
        pre {{ background: #f4f4f4; padding: 1rem; border-radius: 4px; overflow-x: auto; }}
        code {{ background: #f4f4f4; padding: 0.2rem 0.4rem; border-radius: 2px; }}
        blockquote {{ border-left: 4px solid #3498db; margin: 1rem 0; padding-left: 1rem; color: #555; }}
    </style>
</head>
<body>
    <h1>{content.title}</h1>
    <div class="meta">
        {f'<span>Author: {content.author}</span> | ' if content.author else ''}
        {f'<span>Created: {content.created_at.strftime("%Y-%m-%d %H:%M")}</span>' if content.created_at else ''}
    </div>
    <hr>
    <div class="content">
        {html_content}
    </div>
</body>
</html>"""
        return html

    def export_conversation(self, conversation: ConversationExport) -> str:
        """Export conversation to HTML."""
        messages_html = []
        for msg in conversation.messages:
            role = msg.get('role', 'user')
            content = self._markdown_to_html(msg.get('content', ''))
            role_class = 'user' if role == 'user' else 'assistant'
            role_icon = '👤' if role == 'user' else '🤖'

            messages_html.append(f"""
            <div class="message {role_class}">
                <div class="role">{role_icon} {role.title()}</div>
                <div class="content">{content}</div>
            </div>
            """)

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{conversation.title}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 2rem;
            line-height: 1.6;
            color: #333;
            background: #f5f5f5;
        }}
        h1 {{ color: #2c3e50; }}
        .meta {{ color: #7f8c8d; font-size: 0.9rem; margin-bottom: 2rem; }}
        .message {{
            background: white;
            border-radius: 8px;
            padding: 1rem;
            margin-bottom: 1rem;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        }}
        .message.user {{ border-left: 4px solid #3498db; }}
        .message.assistant {{ border-left: 4px solid #27ae60; }}
        .role {{ font-weight: bold; color: #555; margin-bottom: 0.5rem; }}
        pre {{ background: #f4f4f4; padding: 1rem; border-radius: 4px; overflow-x: auto; }}
        code {{ background: #f4f4f4; padding: 0.2rem 0.4rem; border-radius: 2px; }}
    </style>
</head>
<body>
    <h1>{conversation.title}</h1>
    <div class="meta">
        {f'<span>Model: {conversation.model_name}</span> | ' if conversation.model_name else ''}
        {f'<span>Date: {conversation.created_at.strftime("%Y-%m-%d %H:%M")}</span> | ' if conversation.created_at else ''}
        {f'<span>Tokens: {conversation.total_tokens:,}</span>' if conversation.total_tokens else ''}
    </div>
    <div class="conversation">
        {''.join(messages_html)}
    </div>
</body>
</html>"""
        return html

    def _markdown_to_html(self, text: str) -> str:
        """Convert basic Markdown to HTML."""
        if not text:
            return ''

        # Code blocks
        text = re.sub(
            r'```(\w*)\n(.*?)```',
            r'<pre><code class="language-\1">\2</code></pre>',
            text,
            flags=re.DOTALL
        )

        # Inline code
        text = re.sub(r'`([^`]+)`', r'<code>\1</code>', text)

        # Bold
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        text = re.sub(r'__(.+?)__', r'<strong>\1</strong>', text)

        # Italic
        text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
        text = re.sub(r'_(.+?)_', r'<em>\1</em>', text)

        # Headers
        text = re.sub(r'^### (.+)$', r'<h3>\1</h3>', text, flags=re.MULTILINE)
        text = re.sub(r'^## (.+)$', r'<h2>\1</h2>', text, flags=re.MULTILINE)
        text = re.sub(r'^# (.+)$', r'<h1>\1</h1>', text, flags=re.MULTILINE)

        # Links
        text = re.sub(r'\[(.+?)\]\((.+?)\)', r'<a href="\2">\1</a>', text)

        # Lists
        text = re.sub(r'^- (.+)$', r'<li>\1</li>', text, flags=re.MULTILINE)
        text = re.sub(r'(<li>.*</li>\n?)+', r'<ul>\g<0></ul>', text)

        # Blockquotes
        text = re.sub(r'^> (.+)$', r'<blockquote>\1</blockquote>', text, flags=re.MULTILINE)

        # Paragraphs
        paragraphs = text.split('\n\n')
        processed = []
        for p in paragraphs:
            p = p.strip()
            if p and not p.startswith('<'):
                p = f'<p>{p}</p>'
            processed.append(p)
        text = '\n'.join(processed)

        return text


# =============================================================================
# PDF Exporter
# =============================================================================

class PDFExporter:
    """
    Export content to PDF format.

    Requires: reportlab or weasyprint library
    """

    def __init__(self):
        self.available = self._check_availability()

    def _check_availability(self) -> bool:
        """Check if PDF generation is available."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate
            return True
        except ImportError:
            try:
                from weasyprint import HTML
                return True
            except ImportError:
                logger.warning("PDF export requires 'reportlab' or 'weasyprint' package")
                return False

    def export_content(self, content: ExportContent) -> bytes:
        """
        Export content to PDF.

        Args:
            content: Content to export

        Returns:
            PDF file as bytes
        """
        if not self.available:
            raise ImportError("PDF export requires 'reportlab' or 'weasyprint'")

        try:
            return self._export_with_reportlab(content)
        except ImportError:
            return self._export_with_weasyprint(content)

    def _export_with_reportlab(self, content: ExportContent) -> bytes:
        """Generate PDF using ReportLab."""
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=20
        )

        story = []

        # Title
        story.append(Paragraph(content.title, title_style))
        story.append(Spacer(1, 12))

        # Metadata
        if content.author:
            story.append(Paragraph(f"Author: {content.author}", styles['Normal']))
        if content.created_at:
            story.append(Paragraph(
                f"Created: {content.created_at.strftime('%Y-%m-%d %H:%M')}",
                styles['Normal']
            ))

        story.append(Spacer(1, 20))

        # Content - split into paragraphs
        paragraphs = content.content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # Handle basic formatting
                para = para.replace('**', '<b>').replace('*', '<i>')
                story.append(Paragraph(para, styles['Normal']))
                story.append(Spacer(1, 12))

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def _export_with_weasyprint(self, content: ExportContent) -> bytes:
        """Generate PDF using WeasyPrint."""
        from weasyprint import HTML

        html_exporter = HTMLExporter()
        html_content = html_exporter.export_content(content)

        pdf_file = HTML(string=html_content).write_pdf()
        return pdf_file

    def export_conversation(self, conversation: ConversationExport) -> bytes:
        """Export conversation to PDF."""
        if not self.available:
            raise ImportError("PDF export requires 'reportlab' or 'weasyprint'")

        try:
            return self._export_conversation_reportlab(conversation)
        except ImportError:
            return self._export_conversation_weasyprint(conversation)

    def _export_conversation_reportlab(self, conversation: ConversationExport) -> bytes:
        """Generate conversation PDF using ReportLab."""
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)

        styles = getSampleStyleSheet()
        user_style = ParagraphStyle(
            'User',
            parent=styles['Normal'],
            backColor=HexColor('#e3f2fd'),
            borderPadding=10
        )
        assistant_style = ParagraphStyle(
            'Assistant',
            parent=styles['Normal'],
            backColor=HexColor('#e8f5e9'),
            borderPadding=10
        )

        story = []
        story.append(Paragraph(conversation.title, styles['Heading1']))
        story.append(Spacer(1, 20))

        for msg in conversation.messages:
            role = msg.get('role', 'user')
            content = msg.get('content', '')

            style = user_style if role == 'user' else assistant_style
            icon = '👤 User' if role == 'user' else '🤖 Assistant'

            story.append(Paragraph(f"<b>{icon}</b>", styles['Normal']))
            story.append(Spacer(1, 6))
            story.append(Paragraph(content[:2000], style))  # Limit content length
            story.append(Spacer(1, 15))

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def _export_conversation_weasyprint(self, conversation: ConversationExport) -> bytes:
        """Generate conversation PDF using WeasyPrint."""
        from weasyprint import HTML

        html_exporter = HTMLExporter()
        html_content = html_exporter.export_conversation(conversation)
        return HTML(string=html_content).write_pdf()


# =============================================================================
# DOCX Exporter
# =============================================================================

class DOCXExporter:
    """
    Export content to DOCX format.

    Requires: python-docx library
    """

    def __init__(self):
        self.available = self._check_availability()

    def _check_availability(self) -> bool:
        """Check if DOCX generation is available."""
        try:
            from docx import Document
            return True
        except ImportError:
            logger.warning("DOCX export requires 'python-docx' package")
            return False

    def export_content(self, content: ExportContent) -> bytes:
        """
        Export content to DOCX.

        Args:
            content: Content to export

        Returns:
            DOCX file as bytes
        """
        if not self.available:
            raise ImportError("DOCX export requires 'python-docx' package")

        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading(content.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        if content.author or content.created_at:
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if content.author:
                meta_para.add_run(f"Author: {content.author}").italic = True
            if content.author and content.created_at:
                meta_para.add_run(" | ")
            if content.created_at:
                meta_para.add_run(
                    f"Created: {content.created_at.strftime('%Y-%m-%d %H:%M')}"
                ).italic = True

        # Separator
        doc.add_paragraph("_" * 50)

        # Content
        paragraphs = content.content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # Handle headers
                if para.startswith('# '):
                    doc.add_heading(para[2:], level=1)
                elif para.startswith('## '):
                    doc.add_heading(para[3:], level=2)
                elif para.startswith('### '):
                    doc.add_heading(para[4:], level=3)
                else:
                    p = doc.add_paragraph()
                    self._add_formatted_text(p, para)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def export_conversation(self, conversation: ConversationExport) -> bytes:
        """Export conversation to DOCX."""
        if not self.available:
            raise ImportError("DOCX export requires 'python-docx' package")

        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        doc.add_heading(conversation.title, level=0)

        # Metadata
        if conversation.model_name or conversation.created_at:
            meta = doc.add_paragraph()
            if conversation.model_name:
                meta.add_run(f"Model: {conversation.model_name} | ")
            if conversation.created_at:
                meta.add_run(f"Date: {conversation.created_at.strftime('%Y-%m-%d %H:%M')}")
            if conversation.total_tokens:
                meta.add_run(f" | Tokens: {conversation.total_tokens:,}")

        doc.add_paragraph("_" * 50)

        # Messages
        for msg in conversation.messages:
            role = msg.get('role', 'user')
            content = msg.get('content', '')

            # Role header
            role_para = doc.add_paragraph()
            role_run = role_para.add_run(
                f"{'👤 User' if role == 'user' else '🤖 Assistant'}"
            )
            role_run.bold = True
            role_run.font.size = Pt(12)

            # Content
            content_para = doc.add_paragraph()
            self._add_formatted_text(content_para, content)

            # Separator
            doc.add_paragraph("-" * 30)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _add_formatted_text(self, paragraph, text: str):
        """Add formatted text to a paragraph."""
        # Simple formatting - handle bold and italic
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)

        for part in parts:
            if not part:
                continue

            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
            else:
                paragraph.add_run(part)


# =============================================================================
# Export Service (Main Interface)
# =============================================================================

class ExportService:
    """
    Main export service providing unified interface for all export formats.

    Usage:
        from coreapp.services.export_service import export_service

        # Export content
        pdf_bytes = export_service.export_content(content, ExportFormat.PDF)

        # Export conversation
        docx_bytes = export_service.export_conversation(conversation, ExportFormat.DOCX)
    """

    def __init__(self):
        self.markdown = MarkdownExporter()
        self.text = TextExporter()
        self.json = JSONExporter()
        self.html = HTMLExporter()
        self.pdf = PDFExporter()
        self.docx = DOCXExporter()

    def export_content(
        self,
        content: ExportContent,
        format: ExportFormat
    ) -> Union[str, bytes]:
        """
        Export content to specified format.

        Args:
            content: Content to export
            format: Export format

        Returns:
            Exported content (string for text formats, bytes for binary)
        """
        exporters = {
            ExportFormat.MARKDOWN: lambda c: self.markdown.export_content(c),
            ExportFormat.TEXT: lambda c: self.text.export_content(c),
            ExportFormat.JSON: lambda c: self.json.export_content(c),
            ExportFormat.HTML: lambda c: self.html.export_content(c),
            ExportFormat.PDF: lambda c: self.pdf.export_content(c),
            ExportFormat.DOCX: lambda c: self.docx.export_content(c),
        }

        exporter = exporters.get(format)
        if not exporter:
            raise ValueError(f"Unsupported export format: {format}")

        return exporter(content)

    def export_conversation(
        self,
        conversation: ConversationExport,
        format: ExportFormat
    ) -> Union[str, bytes]:
        """
        Export conversation to specified format.

        Args:
            conversation: Conversation to export
            format: Export format

        Returns:
            Exported conversation
        """
        exporters = {
            ExportFormat.MARKDOWN: lambda c: self.markdown.export_conversation(c),
            ExportFormat.TEXT: lambda c: self.text.export_conversation(c),
            ExportFormat.JSON: lambda c: self.json.export_conversation(c),
            ExportFormat.HTML: lambda c: self.html.export_conversation(c),
            ExportFormat.PDF: lambda c: self.pdf.export_conversation(c),
            ExportFormat.DOCX: lambda c: self.docx.export_conversation(c),
        }

        exporter = exporters.get(format)
        if not exporter:
            raise ValueError(f"Unsupported export format: {format}")

        return exporter(conversation)

    def get_content_type(self, format: ExportFormat) -> str:
        """Get MIME content type for format."""
        content_types = {
            ExportFormat.MARKDOWN: 'text/markdown',
            ExportFormat.TEXT: 'text/plain',
            ExportFormat.JSON: 'application/json',
            ExportFormat.HTML: 'text/html',
            ExportFormat.PDF: 'application/pdf',
            ExportFormat.DOCX: 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        }
        return content_types.get(format, 'application/octet-stream')

    def get_file_extension(self, format: ExportFormat) -> str:
        """Get file extension for format."""
        return format.value

    def get_available_formats(self) -> List[Dict[str, Any]]:
        """Get list of available export formats."""
        return [
            {'format': 'md', 'name': 'Markdown', 'available': True},
            {'format': 'txt', 'name': 'Plain Text', 'available': True},
            {'format': 'json', 'name': 'JSON', 'available': True},
            {'format': 'html', 'name': 'HTML', 'available': True},
            {'format': 'pdf', 'name': 'PDF', 'available': self.pdf.available},
            {'format': 'docx', 'name': 'Word Document', 'available': self.docx.available},
        ]


# =============================================================================
# Singleton Instance
# =============================================================================

export_service = ExportService()
