"""
Export service for MultinotesAI.

This module provides functionality to export notes and content to various formats:
- PDF
- DOCX (Microsoft Word)
- Markdown
- HTML
- Plain Text
"""

import io
import os
import logging
from datetime import datetime
from typing import Optional, List, Dict, Any

from django.conf import settings
from django.template.loader import render_to_string
from django.utils.html import strip_tags

logger = logging.getLogger(__name__)


# =============================================================================
# Export Configuration
# =============================================================================

class ExportConfig:
    """Configuration for export operations."""

    # PDF settings
    PDF_PAGE_SIZE = 'A4'
    PDF_MARGIN_TOP = '2cm'
    PDF_MARGIN_BOTTOM = '2cm'
    PDF_MARGIN_LEFT = '2.5cm'
    PDF_MARGIN_RIGHT = '2.5cm'

    # DOCX settings
    DOCX_FONT_NAME = 'Calibri'
    DOCX_FONT_SIZE = 11
    DOCX_HEADING_SIZE = 16

    # Export limits
    MAX_NOTES_PER_EXPORT = 100
    MAX_EXPORT_SIZE_MB = 50


# =============================================================================
# Base Exporter Class
# =============================================================================

class BaseExporter:
    """Base class for all exporters."""

    def __init__(self, notes: List[Dict], options: Optional[Dict] = None):
        """
        Initialize exporter.

        Args:
            notes: List of note dictionaries with title, content, etc.
            options: Export options (format-specific)
        """
        self.notes = notes
        self.options = options or {}

    def export(self) -> bytes:
        """
        Export notes to the target format.

        Returns:
            bytes: Exported content
        """
        raise NotImplementedError("Subclasses must implement export()")

    def get_filename(self, base_name: str = "export") -> str:
        """Generate filename with timestamp."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{base_name}_{timestamp}.{self.file_extension}"

    @property
    def file_extension(self) -> str:
        """Return the file extension for this format."""
        raise NotImplementedError

    @property
    def content_type(self) -> str:
        """Return the MIME content type for this format."""
        raise NotImplementedError


# =============================================================================
# Markdown Exporter
# =============================================================================

class MarkdownExporter(BaseExporter):
    """Export notes to Markdown format."""

    file_extension = "md"
    content_type = "text/markdown"

    def export(self) -> bytes:
        """Export notes to Markdown."""
        lines = []

        for note in self.notes:
            # Title
            title = note.get('title', 'Untitled')
            lines.append(f"# {title}\n")

            # Metadata
            if self.options.get('include_metadata', True):
                created_at = note.get('created_at', '')
                updated_at = note.get('updated_at', '')
                if created_at:
                    lines.append(f"*Created: {created_at}*\n")
                if updated_at:
                    lines.append(f"*Updated: {updated_at}*\n")
                lines.append("\n")

            # Content
            content = note.get('content', '')
            lines.append(content)
            lines.append("\n\n---\n\n")

        result = "\n".join(lines)
        return result.encode('utf-8')


# =============================================================================
# HTML Exporter
# =============================================================================

class HTMLExporter(BaseExporter):
    """Export notes to HTML format."""

    file_extension = "html"
    content_type = "text/html"

    def export(self) -> bytes:
        """Export notes to HTML."""
        template = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>MultinotesAI Export</title>
    <style>
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
            color: #333;
        }
        .note {
            margin-bottom: 40px;
            padding-bottom: 20px;
            border-bottom: 1px solid #eee;
        }
        .note h1 {
            color: #2c3e50;
            margin-bottom: 10px;
        }
        .note-meta {
            color: #666;
            font-size: 0.9em;
            margin-bottom: 15px;
        }
        .note-content {
            white-space: pre-wrap;
        }
        .header {
            text-align: center;
            margin-bottom: 40px;
            padding-bottom: 20px;
            border-bottom: 2px solid #3498db;
        }
        .footer {
            text-align: center;
            color: #666;
            font-size: 0.8em;
            margin-top: 40px;
        }
    </style>
</head>
<body>
    <div class="header">
        <h1>MultinotesAI Export</h1>
        <p>Exported on {export_date}</p>
    </div>

    {notes_html}

    <div class="footer">
        <p>Exported from MultinotesAI</p>
    </div>
</body>
</html>
"""

        notes_html = []
        for note in self.notes:
            title = note.get('title', 'Untitled')
            content = note.get('content', '')
            created_at = note.get('created_at', '')
            updated_at = note.get('updated_at', '')

            meta_parts = []
            if created_at:
                meta_parts.append(f"Created: {created_at}")
            if updated_at:
                meta_parts.append(f"Updated: {updated_at}")

            note_html = f"""
    <div class="note">
        <h1>{title}</h1>
        <div class="note-meta">{' | '.join(meta_parts)}</div>
        <div class="note-content">{content}</div>
    </div>
"""
            notes_html.append(note_html)

        result = template.format(
            export_date=datetime.now().strftime("%Y-%m-%d %H:%M"),
            notes_html="\n".join(notes_html)
        )

        return result.encode('utf-8')


# =============================================================================
# Plain Text Exporter
# =============================================================================

class PlainTextExporter(BaseExporter):
    """Export notes to plain text format."""

    file_extension = "txt"
    content_type = "text/plain"

    def export(self) -> bytes:
        """Export notes to plain text."""
        lines = []
        separator = "=" * 60

        lines.append("MultinotesAI Export")
        lines.append(f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        lines.append(separator)
        lines.append("")

        for i, note in enumerate(self.notes, 1):
            title = note.get('title', 'Untitled')
            content = strip_tags(note.get('content', ''))

            lines.append(f"Note {i}: {title}")
            lines.append("-" * 40)

            if self.options.get('include_metadata', True):
                created_at = note.get('created_at', '')
                if created_at:
                    lines.append(f"Created: {created_at}")

            lines.append("")
            lines.append(content)
            lines.append("")
            lines.append(separator)
            lines.append("")

        result = "\n".join(lines)
        return result.encode('utf-8')


# =============================================================================
# PDF Exporter
# =============================================================================

class PDFExporter(BaseExporter):
    """Export notes to PDF format using WeasyPrint."""

    file_extension = "pdf"
    content_type = "application/pdf"

    def export(self) -> bytes:
        """Export notes to PDF."""
        try:
            from weasyprint import HTML, CSS
        except ImportError:
            logger.error("WeasyPrint not installed. Please install it for PDF export.")
            raise ImportError("PDF export requires WeasyPrint: pip install weasyprint")

        # First generate HTML
        html_exporter = HTMLExporter(self.notes, self.options)
        html_content = html_exporter.export().decode('utf-8')

        # Add print-specific CSS
        print_css = CSS(string="""
            @page {
                size: A4;
                margin: 2cm;
            }
            body {
                font-size: 12pt;
            }
            .note {
                page-break-inside: avoid;
            }
        """)

        # Generate PDF
        html = HTML(string=html_content)
        pdf_bytes = html.write_pdf(stylesheets=[print_css])

        return pdf_bytes


# =============================================================================
# DOCX Exporter
# =============================================================================

class DOCXExporter(BaseExporter):
    """Export notes to Microsoft Word DOCX format."""

    file_extension = "docx"
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def export(self) -> bytes:
        """Export notes to DOCX."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx not installed. Please install it for DOCX export.")
            raise ImportError("DOCX export requires python-docx: pip install python-docx")

        # Create document
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        style.font.name = ExportConfig.DOCX_FONT_NAME
        style.font.size = Pt(ExportConfig.DOCX_FONT_SIZE)

        # Add title
        title = doc.add_heading('MultinotesAI Export', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add export date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        date_run.italic = True

        doc.add_paragraph()  # Spacing

        # Add notes
        for note in self.notes:
            title = note.get('title', 'Untitled')
            content = strip_tags(note.get('content', ''))

            # Note title
            doc.add_heading(title, level=1)

            # Metadata
            if self.options.get('include_metadata', True):
                created_at = note.get('created_at', '')
                if created_at:
                    meta_para = doc.add_paragraph()
                    meta_run = meta_para.add_run(f"Created: {created_at}")
                    meta_run.italic = True
                    meta_run.font.size = Pt(10)

            # Content
            doc.add_paragraph(content)

            # Separator
            doc.add_paragraph("_" * 50)
            doc.add_paragraph()

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.read()


# =============================================================================
# Export Service
# =============================================================================

class ExportService:
    """
    Main export service for generating exports in various formats.

    Usage:
        service = ExportService()
        pdf_content, filename, content_type = service.export_notes(
            notes=[{'title': 'Note 1', 'content': 'Content...'}],
            format='pdf'
        )
    """

    EXPORTERS = {
        'md': MarkdownExporter,
        'markdown': MarkdownExporter,
        'html': HTMLExporter,
        'txt': PlainTextExporter,
        'text': PlainTextExporter,
        'pdf': PDFExporter,
        'docx': DOCXExporter,
        'word': DOCXExporter,
    }

    SUPPORTED_FORMATS = list(EXPORTERS.keys())

    def export_notes(
        self,
        notes: List[Dict],
        format: str = 'pdf',
        options: Optional[Dict] = None
    ) -> tuple:
        """
        Export notes to specified format.

        Args:
            notes: List of note dictionaries
            format: Export format (pdf, docx, md, html, txt)
            options: Export options

        Returns:
            tuple: (content bytes, filename, content_type)
        """
        format = format.lower()

        if format not in self.EXPORTERS:
            raise ValueError(f"Unsupported format: {format}. Supported: {self.SUPPORTED_FORMATS}")

        if len(notes) > ExportConfig.MAX_NOTES_PER_EXPORT:
            raise ValueError(f"Maximum {ExportConfig.MAX_NOTES_PER_EXPORT} notes per export")

        exporter_class = self.EXPORTERS[format]
        exporter = exporter_class(notes, options)

        content = exporter.export()
        filename = exporter.get_filename("notes_export")
        content_type = exporter.content_type

        return content, filename, content_type

    def export_single_note(
        self,
        note: Dict,
        format: str = 'pdf',
        options: Optional[Dict] = None
    ) -> tuple:
        """
        Export a single note.

        Args:
            note: Note dictionary
            format: Export format
            options: Export options

        Returns:
            tuple: (content bytes, filename, content_type)
        """
        return self.export_notes([note], format, options)

    @classmethod
    def get_supported_formats(cls) -> List[str]:
        """Get list of supported export formats."""
        return list(set([
            'pdf', 'docx', 'md', 'html', 'txt'
        ]))


# =============================================================================
# Singleton Instance
# =============================================================================

export_service = ExportService()
